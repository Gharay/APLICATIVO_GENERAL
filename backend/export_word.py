"""
backend/export_word.py
Informe de Gestión Catastral — Dirección de Acceso a Tierras
Estructura basada en plantilla institucional ANT.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io

# ============================================================
# CONSTANTES DE ESTILO
# ============================================================
FUENTE        = 'Aptos Narrow'
NEGRO         = RGBColor(0x00, 0x00, 0x00)
COLOR_HEADER  = '92D050'   # Verde ANT — encabezados de tabla
COLOR_ALTERNO = 'DAF2D0'   # Verde claro — filas alternas
LANG          = 'es-CO'    # Idioma predeterminado del documento

# ============================================================
# UTILIDADES
# ============================================================
def _v(datos, clave, defecto='—'):
    """Retorna el valor del dict como string, o defecto si vacío/nulo."""
    if not datos:
        return defecto
    v = datos.get(clave)
    return str(v) if v not in (None, '', 'None') else defecto

def _bool_s(valor):
    return 'Sí' if valor else 'No'

def _aplicar_fuente(run, bold=False, italic=False):
    run.font.name      = FUENTE
    run.font.size      = Pt(11)
    run.font.color.rgb = NEGRO
    run.font.bold      = bold
    run.font.italic    = italic
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:lang')):
        rPr.remove(existing)
    lang_el = OxmlElement('w:lang')
    lang_el.set(qn('w:val'),      LANG)
    lang_el.set(qn('w:eastAsia'), 'zxx')
    lang_el.set(qn('w:bidi'),     'zxx')
    rPr.append(lang_el)

def _aplicar_interlineado(p):
    """Aplica interlineado 1,5 al párrafo."""
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def _p(doc, texto, bold=False, italic=False, align=None):
    p   = doc.add_paragraph()
    run = p.add_run(str(texto))
    _aplicar_fuente(run, bold=bold, italic=italic)
    _aplicar_interlineado(p)
    if align:
        p.alignment = align
    return p

def _shd(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'),  hex_fill)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),   'clear')
    tcPr.append(shd)

def _heading(doc, texto, nivel):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        _aplicar_fuente(run, bold=True)
    _aplicar_interlineado(h)
    return h

def _caption(doc, texto):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    _aplicar_fuente(run, italic=True)
    _aplicar_interlineado(p)
    return p

def _placeholder(doc, texto='[COMPLETAR]'):
    p   = doc.add_paragraph()
    run = p.add_run(texto)
    _aplicar_fuente(run, italic=True)
    _aplicar_interlineado(p)
    return p

def _p_bold(doc, *partes):
    """Párrafo con texto mixto normal/negrita.
    Uso: _p_bold(doc, ("texto ", False), ("VALOR", True))
    """
    p = doc.add_paragraph()
    _aplicar_interlineado(p)
    for texto, bold in partes:
        run = p.add_run(str(texto) if texto is not None else '—')
        _aplicar_fuente(run, bold=bold)
    return p

# ── Escritura de celda (evita runs vacíos) ──────────────────
def _set_cell(cell, texto, bold=False, italic=False, shd_color=None):
    """Limpia la celda y escribe texto con estilo."""
    for par in cell.paragraphs:
        for run in par.runs:
            run.text = ''
    cell.text = ''
    run = cell.paragraphs[0].add_run(str(texto) if texto not in (None, '', 'None') else '—')
    run.font.name      = FUENTE
    run.font.size      = Pt(9)
    run.font.color.rgb = NEGRO
    run.font.bold      = bold
    run.font.italic    = italic
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:lang')):
        rPr.remove(existing)
    lang_el = OxmlElement('w:lang')
    lang_el.set(qn('w:val'),      LANG)
    lang_el.set(qn('w:eastAsia'), 'zxx')
    lang_el.set(qn('w:bidi'),     'zxx')
    rPr.append(lang_el)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if shd_color:
        _shd(cell, shd_color)

# ── Tabla 2 columnas ────────────────────────────────────────
def _tabla_dos_col(doc, filas, h0='CONCEPTO', h1='VALOR'):
    tabla = doc.add_table(rows=len(filas) + 1, cols=2)
    tabla.style = 'Table Grid'
    tabla.columns[0].width = Cm(8)
    tabla.columns[1].width = Cm(9)

    enc = tabla.rows[0].cells
    _set_cell(enc[0], h0, bold=True, shd_color=COLOR_HEADER)
    _set_cell(enc[1], h1, bold=True, shd_color=COLOR_HEADER)

    for idx, (concepto, valor) in enumerate(filas, start=1):
        fila  = tabla.rows[idx].cells
        color = COLOR_ALTERNO if idx % 2 == 0 else None
        _set_cell(fila[0], concepto, bold=True,  shd_color=color)
        _set_cell(fila[1], valor,    bold=False, shd_color=color)

    return tabla

# ── Tabla de colindantes ─────────────────────────────────────
def _tabla_colindantes(doc, lista, titulo_tabla=None):
    """Tabla completa de colindantes con todos los campos capturados."""
    if not lista:
        _p(doc, 'No hay colindantes registrados para este tipo.', italic=True)
        return

    _p(doc, f'Total de colindantes encontrados: {len(lista)}')

    headers = [
        'N°', 'Cardinalidad', 'FMI', 'N° Predial',
        'Nombre Predio', 'Doc. Origen',
        'Prop. Inicial', 'Tipo'
    ]
    anchos = [
        Cm(0.94), Cm(2.25), Cm(1.5), Cm(2.25), Cm(2.25), Cm(4.25), Cm(3), Cm(2.25)
    ]

    tabla = doc.add_table(rows=len(lista) + 1, cols=len(headers))
    tabla.style = 'Table Grid'

    for i, ancho in enumerate(anchos):
        for cell in tabla.columns[i].cells:
            cell.width = ancho

    for i, h in enumerate(headers):
        _set_cell(tabla.rows[0].cells[i], h, bold=True, shd_color=COLOR_HEADER)

    for r_idx, col in enumerate(lista, start=1):
        fila  = tabla.rows[r_idx].cells
        color = COLOR_ALTERNO if r_idx % 2 == 0 else None

        valores = [
            str(col.get('numero_colindante', r_idx)),
            _v(col, 'cardinalidad'),
            _v(col, 'fmi_colindante'),
            _v(col, 'numero_predial_colindante'),
            _v(col, 'nombre_predio_colindante'),
            _v(col, 'documento_origen_colindante'),
            _v(col, 'propietario_inicial_colindante'),
            _v(col, 'tipo_colindante'),
        ]

        for c_idx, valor in enumerate(valores):
            _set_cell(fila[c_idx], valor, shd_color=color)

    if titulo_tabla:
        _caption(doc, titulo_tabla)

# --- Tabla de colindantes resumen

def _tabla_colindantes_resumen_registrales(doc, lista, titulo_tabla=None):
    """Tabla completa de colindantes final."""
    if not lista:
        _p(doc, 'No hay colindantes registrados para este tipo.', italic=True)
        return

    _p(doc, f'Total de colindantes encontrados: {len(lista)}')

    headers = [
        'N°', 'Cardinalidad en Resolución', 'Colindante en Resolución de Adjudicación', '¿Colindante Verificado?', 'Colindante Final', 'FMI', 'Cardinalidad Final'
    ]
    anchos = [
        Cm(0.8), Cm(2.5), Cm(3.8), Cm(2.5), Cm(4), Cm(2.3), Cm(2.5)
    ]

    tabla = doc.add_table(rows=len(lista) + 1, cols=len(headers))
    tabla.style = 'Table Grid'

    for i, ancho in enumerate(anchos):
        for cell in tabla.columns[i].cells:
            cell.width = ancho

    for i, h in enumerate(headers):
        _set_cell(tabla.rows[0].cells[i], h, bold=True, shd_color=COLOR_HEADER)

    for r_idx, col in enumerate(lista, start=1):
        fila  = tabla.rows[r_idx].cells
        color = COLOR_ALTERNO if r_idx % 2 == 0 else None

        valores = [
            str(col.get('numero_colindante', r_idx)),
            _v(col, 'cardinalidad'),
            _v(col, 'propietario_inicial_colindante'),
            _v(col, 'estado_validacion'),
            _v(col, 'propietario_actual_vur_colindante'),
            _v(col, 'fmi_colindante'),
            _v(col, 'cardinalidad'),
        ]

        for c_idx, valor in enumerate(valores):
            _set_cell(fila[c_idx], valor, shd_color=color)

    if titulo_tabla:
        _caption(doc, titulo_tabla)

# ── Tabla de colindantes ─────────────────────────────────────
def _tabla_colindantes_final(doc, lista, titulo_tabla=None):
    """Tabla completa de colindantes final."""
    if not lista:
        _p(doc, 'No hay colindantes registrados para este tipo.', italic=True)
        return

    _p(doc, f'Total de colindantes encontrados: {len(lista)}')

    headers = [
        'N°', 'FMI', 'Colindante Final', '¿Colindante Verificado?', '¿Necesario Validar en Campo?', 'Nombre del Predio', 'Observación'
    ]
    anchos = [
        Cm(0.8), Cm(1.8), Cm(4), Cm(2.6), Cm(1.9), Cm(2.3), Cm(5)
    ]

    tabla = doc.add_table(rows=len(lista) + 1, cols=len(headers))
    tabla.style = 'Table Grid'

    for i, ancho in enumerate(anchos):
        for cell in tabla.columns[i].cells:
            cell.width = ancho

    for i, h in enumerate(headers):
        _set_cell(tabla.rows[0].cells[i], h, bold=True, shd_color=COLOR_HEADER)

    for r_idx, col in enumerate(lista, start=1):
        fila  = tabla.rows[r_idx].cells
        color = COLOR_ALTERNO if r_idx % 2 == 0 else None

        valores = [
            str(col.get('numero_colindante', r_idx)),
            _v(col, 'fmi_colindante'),
            _v(col, 'propietario_actual_vur_colindante'),
            _v(col, 'estado_validacion'),
            _v(col, 'requiere_verificacion_campo'),
            _v(col, 'nombre_predio_colindante'),
            'Se identifica la colindancia en ' + _v(col, 'tipo_colindante'),
        ]

        for c_idx, valor in enumerate(valores):
            _set_cell(fila[c_idx], valor, shd_color=color)

    if titulo_tabla:
        _caption(doc, titulo_tabla)


# ── Mini-tabla resumen de validación de colindantes ──────────
def _tabla_resumen_colindantes(doc, colindantes):
    validados = sum(1 for c in colindantes
                    if (c.get('estado_validacion') or '').upper() == 'VALIDADO')
    no_verif  = sum(1 for c in colindantes
                    if (c.get('estado_validacion') or '').upper() == 'NO_VERIFICADO')
    baldios   = sum(1 for c in colindantes if c.get('es_presunto_baldio'))
    req_campo = sum(1 for c in colindantes if c.get('requiere_verificacion_campo'))

    _tabla_dos_col(doc, [
        ('Total de colindantes encontrados',            str(len(colindantes))),
        ('Colindantes validados',           str(validados)),
        ('Colindantes no verificados',      str(no_verif)),
        ('Requieren verificación en campo', str(req_campo)),
        ('Presuntos baldíos',               str(baldios)),
    ], h0='INDICADOR', h1='CANTIDAD')

def _bloques_colindantes(doc, lista, ilustracion_inicio):
    """
    Genera para cada colindante: párrafo descriptivo + caption de ilustración.
    Retorna el siguiente número de ilustración disponible.
    """
    num_ilus = ilustracion_inicio

    for col in lista:
        n        = str(col.get('numero_colindante') or '—')
        card     = _v(col, 'cardinalidad')
        fmi_col  = _v(col, 'fmi_colindante')
        id_ant   = _v(col, 'id_ant')
        nombre   = _v(col, 'nombre_predio_colindante')
        doc_orig = _v(col, 'documento_origen_colindante')
        num_pred = _v(col, 'numero_predial_colindante')
        nom_p_inicial = _v(col, 'propietario_inicial_colindante')
        nom_p_actual = _v(col, 'propietario_actual_vur_colindante')
        fuente_validacion = _v(col, 'fuente_validacion')

        _p_bold(doc,
            (f'Lindero {n} al ', False),
            (card, True),
            (': FMI: ', False),
            (fmi_col, True),
            (', ', False),
            (id_ant, True),
            (' Predio denominado ', False),
            (nombre, True),
            ('. Según consulta VUR el predio tiene origen mediante la ', False),
            (doc_orig, True),
            ('. Según base catastral se identifica que le corresponde el número predial: ', False),
            (num_pred, True),
            ('. La información jurídica señala a ', False),
            (nom_p_inicial, True),
            (' como adjudicatario inicial, asimismo se identifica a ', False),
            (nom_p_actual, True),
            (' como propietario actual del predio colindante.', False)
        )

        _p_bold(doc,
            ('Fuente de validación: ', False),
            (fuente_validacion, False)
        )

        _p_bold(doc,
            ('El predio tuvo un proceso de segregación desde el predio matriz identificado con FMI: ', False),
            (_v(col, 'id_predio_matriz_colindante'), True), ('.', False),
        )

        _p_bold(doc,
            ('El predio ha generado los siguientes predios derivados, identificados con FMI: ', False),
            (_v(col, 'derivadas_ids_colindante'), True), ('.', False),
        )

        _p_bold(doc,
            ('Según la revisión jurídica del FMI, el predio no registra Folio Matriz ni presenta algún proceso de desenglobe o venta parcial, motivo por el cual no se identifican FMI segregados y se determina que su extensión jurídica sigue siendo la misma con la que se originó su apertura ', False),
        )

        _p_bold(doc,
            ('Se realiza el proceso de reconstrucción de los linderos a partir de la descripción de rumbos y distancias definidas en el documento que da origen al predio colindante.', False),
        )
        _caption(doc, '')
        _caption(doc, f'Ilustración {num_ilus}. Colindante {n} - {card} - {nombre}')
        num_ilus += 1

    return num_ilus

def _tabla_evaluacion_final_resumen(doc):
    """Tabla 3×4: VALIDACIÓN / VERIFICACIÓN / OBSERVACIÓN.
    Fila header + filas FORMA, UBICACIÓN, COLINDANTES."""
    tabla = doc.add_table(rows=4, cols=3)
    tabla.style = 'Table Grid'

    ancho_col = Cm(5)
    for col in tabla.columns:
        for cell in col.cells:
            cell.width = ancho_col

    # Fila 0 — encabezados verdes
    enc = tabla.rows[0].cells
    _set_cell(enc[0], 'VALIDACIÓN',   bold=True, shd_color=COLOR_HEADER)
    _set_cell(enc[1], 'VERIFICACIÓN', bold=True, shd_color=COLOR_HEADER)
    _set_cell(enc[2], 'OBSERVACIÓN',  bold=True, shd_color=COLOR_HEADER)

    # Filas 1-3: etiqueta en col-0, celdas 1 y 2 vacías
    for i, etiqueta in enumerate(['FORMA', 'UBICACIÓN', 'COLINDANTES'], start=1):
        fila = tabla.rows[i].cells
        _set_cell(fila[0], etiqueta, bold=True)
        for j in (1, 2):
            fila[j].text = ''   # celda vacía (sin guion)

# ============================================================
# FUNCIÓN PRINCIPAL
# ============================================================
def exportar_word(datos_predio, upload_folder=None):
    doc = Document()

    # ── Idioma predeterminado: español Colombia ─────────────────
    styles_el    = doc.styles.element
    doc_defaults = styles_el.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_el.insert(0, doc_defaults)
    rPrDef = doc_defaults.find(qn('w:rPrDefault'))
    if rPrDef is None:
        rPrDef = OxmlElement('w:rPrDefault')
        doc_defaults.append(rPrDef)
    rPr_d = rPrDef.find(qn('w:rPr'))
    if rPr_d is None:
        rPr_d = OxmlElement('w:rPr')
        rPrDef.append(rPr_d)
    for ex in rPr_d.findall(qn('w:lang')):
        rPr_d.remove(ex)
    lang_def = OxmlElement('w:lang')
    lang_def.set(qn('w:val'),      LANG)
    lang_def.set(qn('w:eastAsia'), 'zxx')
    lang_def.set(qn('w:bidi'),     'zxx')
    rPr_d.append(lang_def)

    # ── Interlineado 1,5 en estilo Normal ───────────────────────
    doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Acceso defensivo a cada sección
    predio       = datos_predio.get('predio')       or {}
    localizacion = datos_predio.get('localizacion') or {}
    catastral    = datos_predio.get('catastral')    or {}
    juridica     = datos_predio.get('juridica')     or {}
    colindantes  = datos_predio.get('colindantes')  or []

    col_registral     = [c for c in colindantes
                         if 'REGISTRAL' in (c.get('tipo_colindante') or '').upper()]
    col_catastral_lst = [c for c in colindantes
                         if 'CATASTRAL' in (c.get('tipo_colindante') or '').upper()]

    nombre_actual = _v(juridica, 'nombre_predio_actual',
                       _v(predio, 'nombre_predio_vur'))
    num_predial   = _v(catastral, 'numero_predial')
    fmi           = _v(predio, 'fmi')
    dpto          = _v(localizacion, 'dpto_espacial')
    mpio          = _v(localizacion, 'municipio_espacial')
    vereda        = _v(localizacion, 'vereda_espacial')
    documento_origen_inmueble = _v(juridica, 'documento_origen_inmueble')
    documento_ultima_anotacion = _v(juridica, 'documento_ultima_anotacion')

    # ═══════════════════════════════════════════════════════
    # ENCABEZADO
    # ═══════════════════════════════════════════════════════
    _p(doc, 'INFORME DE GESTIÓN CATASTRAL', bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, 'DIRECCIÓN DE ACCESO A TIERRAS',
       align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════════════════════════════════════════════════
    # 1. EVALUACIÓN INICIAL
    # ═══════════════════════════════════════════════════════
    _heading(doc, 'EVALUACIÓN INICIAL', 1)

    _tabla_dos_col(doc, [
        ('ID del Predio',                             _v(predio, 'id')),
        ('FMI',                                       fmi),
        ('Nombre del predio (VUR)',                   _v(predio, 'nombre_predio_vur')),
        ('Propietario (VUR)',                         _v(predio, 'propietario_vur')),
        ('# Documento del Propietario (VUR)',         _v(predio, 'numero_documento_propietario_vur')),
        ('Área según VUR (ha)',                       _v(predio, 'area_vur', 'N/A')),
        ('¿Nace de un predio matriz?',                _bool_s(predio.get('es_derivado'))),
        ('ID Predio Matriz',                          _v(predio, 'id_predio_matriz', 'N/A')),
        ('¿Tiene predios derivados?',                 _bool_s(predio.get('tiene_derivadas'))),
        ('IDs Predios Derivados',                     _v(predio, 'derivadas_ids', 'N/A')),
        ('Proceso englobe / desenglobe',              _bool_s(predio.get('proceso_englobe_desenglobe'))),
        ('Estado',                                    _v(predio, 'estado')),
    ], h0='CONCEPTO', h1='ATRIBUTO')
    _caption(doc, 'Tabla 1. Evaluación inicial del predio')

    _p_bold(doc,
        ('El predio objeto de estudio denominado ', False),
        (nombre_actual, True),
        (', identificado con número predial ', False),
        (num_predial, True),
        (' y Folio de Matrícula Inmobiliaria (FMI) ', False),
        (fmi, True),
        ('. Se encuentra ubicado en el departamento de ', False),
        (dpto, True),
        (', municipio de ', False),
        (mpio, True),
        (' en la vereda ', False),
        (vereda, True),
        ('. Registra un área jurídica de ', False),
        (_v(juridica, 'area_juridica_ha', 'N/A'), True),
        (' ha de acuerdo con la ', False),
        (_v(juridica, 'documento_origen_inmueble'), True),
        (', documento por el cual se da apertura al FMI.', False),
    )

    if predio.get('es_derivado') and predio.get('id_predio_matriz'):
        _p_bold(doc,
            ('El predio tuvo un proceso de segregación desde el predio matriz identificado con el FMI: ', False),
            (_v(predio, 'id_predio_matriz'), True), ('.', False),
        )

    if predio.get('tiene_derivadas') and predio.get('derivadas_ids'):
        _p_bold(doc,
            ('El predio ha generado los siguientes predios derivados, identificados con FMI: ', False),
            (_v(predio, 'derivadas_ids'), True), ('.', False),
        )

    _caption(doc, '')
    _caption(doc, f"Ilustración 1. Tradición {fmi} — {_v(predio, 'id')}")

    _p_bold(doc,
        ('Según la revisión jurídica del FMI, el predio no registra Folio Matriz ni presenta algún proceso de desenglobe o venta parcial, motivo por el cual no se identifican FMI segregados y se determina que su extensión jurídica sigue siendo la misma con la que se originó su apertura. Por lo cual se solicita la resolución de adjudicación para validar su ubicación y correlación con el polígono identificado en la malla catastral. Cabe resaltar que este no cuenta con levantamiento predial realizado por la Agencia Nacional de Tierras. ', False),
    )

    # ═══════════════════════════════════════════════════════
    # 2. DESARROLLO DEL ANÁLISIS
    # ═══════════════════════════════════════════════════════
    _heading(doc, 'DESARROLLO DEL ANÁLISIS', 1)

    # ── 2.1 Análisis jurídico ────────────────────────────
    _heading(doc, 'ANÁLISIS JURÍDICO', 2)

    _tabla_dos_col(doc, [
        ('Nombre del predio actual',             _v(juridica, 'nombre_predio_actual')),
        ('Nombre del predio en doc. de origen',  _v(juridica, 'nombre_predio_origen')),
        ('Documento jurídico de origen',         _v(juridica, 'documento_origen_inmueble')),
        ('Área jurídica (ha)',                   _v(juridica, 'area_juridica_ha', 'N/A')),
        ('Adjudicatario inicial',                _v(juridica, 'adjudicatario_inicial')),
        ('Tipo de instrumento',                  _v(juridica, 'tipo_instrumento')),
        ('Fecha de instrumento',                 _v(juridica, 'fecha_instrumento')),
        ('Tipo de predio',                       _v(juridica, 'tipo_predio')),
        ('Estado del folio',                     _v(juridica, 'estado_folio')),
        ('Fecha apertura del folio',             _v(juridica, 'fecha_apertura_folio')),
        ('Documento primera anotación',          _v(juridica, 'documento_primera_anotacion')),
        ('Documento última anotación (Propietario Actual) ',           _v(juridica, 'documento_ultima_anotacion')),
        ('Complementaciones',                    _v(juridica, 'complementaciones')),
        ('Cabida y linderos',                    _v(juridica, 'cabida_linderos')),
        ('Salvedades',                           _v(juridica, 'salvedades')),
    ])

    _caption(doc, 'Tabla 2. Análisis jurídico del predio')

    _p_bold(doc,
        ('El predio objeto de estudio nace a través de la ', False),
        (_v(juridica, 'documento_origen_inmueble'), True),
        (', en la cual se adjudica el derecho de dominio de un predio rural denominado ', False),
        (_v(juridica, 'nombre_predio_origen'), True),
        (', con una extensión superficial de ', False),
        (_v(juridica, 'area_juridica_ha', 'N/A'), True),
        (' ha. Fue adjudicado a ', False),
        (_v(juridica, 'adjudicatario_inicial'), True),
        ('.', False),
    )

    _p_bold(doc,
        ('Se realizo el proceso de reconstrucción y georreferenciación del plano, así como la revisión de la descripción de los colindantes en el documento que da origen al predio objeto de estudio.', False),
    )
    _caption(doc, '')
    _caption(doc, f"Ilustración 2. Plano {documento_origen_inmueble}")

    _p_bold(doc,
        ('El último registro según la cadena de tradición hace referencia a la ', False),
        (_v(juridica, 'documento_ultima_anotacion'), True),
        ('. Siendo la Escritura XXXX, el documento que define al último propietario del predio. No obstante se registra, en la anotación X, la suspensión provisional a la libre disposición de dominio en proceso de Justicia y Paz, mismo que a la fecha se encuentra....', False),
    )

    _p_bold(doc,
        ('Lo anteriormente expuesto evidencia que existen títulos debidamente inscritos y otorgados antes de la entrada en vigencia de Ley 160 de 1994, se ilustran las constancias de tradiciones de dominio. Por lo tanto se puede acreditar, en base al artículo 48 de la Ley 160, que la naturaleza jurídica del presente predio es PRIVADA.', False),
    )
    _caption(doc, '')
    _p_bold(doc,
        ('Se cuenta con levantamiento topográfico realizado por la Agencia Nacional de Tierras, Dirección de Acceso a Tierras - DAT, en XX de XXXX en el cual se levanto...', False),
    )
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Levantamiento topográfico XXXXX')

    # ── 2.2 Análisis catastral ───────────────────────────
    _heading(doc, 'ANÁLISIS CATASTRAL', 2)

    _caption(doc, '')
    _caption(doc, f"Ilustración 3. Información catastral del predio objeto de estudio {fmi}")

    _tabla_dos_col(doc, [
        ('Número predial',                   _v(catastral, 'numero_predial')),
        ('FMI en R1/R2',                     _v(catastral, 'fmi_r1_r2')),
        ('Nombre del predio en R1/R2',       _v(catastral, 'nombre_predio_r1_r2')),
        ('Propietario (R1/R2)',              _v(catastral, 'propietario')),
        ('Cédula del propietario',           _v(catastral, 'cedula_propietario')),
        ('Área de terreno — polígono (ha)',  _v(catastral, 'area_terreno_ha', 'N/A')),
        ('Área de terreno — R1/R2 (ha)',     _v(catastral, 'area_terreno_ha_r1r2', 'N/A')),
        ('Destino económico',                _v(catastral, 'destino_economico')),
    ])

    _caption(doc, 'Tabla 3. Información catastral del predio')
    _caption(doc, '')
    _caption(doc, 'Ilustración 4. Identificación catastral del predio objeto de estudio y sus colindantes')
    _caption(doc, '')
    _caption(doc, 'Ilustración 5. Comparativa entre polígono catastral e información jurídica del predio objeto de estudio')

    _p_bold(doc,
        ('Como se puede ver en la imagen anterior, el polígono definido en la resolución de adjudicación del predio objeto de estudio no coincide con el polígono catastral vigente. Por lo tanto, para dar inicio al estudio del predio se realiza un ANÁLISIS TÉCNICO que incluye la revisión de los cambios en las coberturas a partir de un estudio multitemporal con el objetivo de observar los cambios en las formas del predio y la existencia de linderos arcifinios y/o artificiales. Asimismo, un ANÁLISIS DE COLINDANTES identificados en los insumos catastrales y los insumos registrales, posteriormente se validan las coincidencias entre ambos insumos y se realiza el análisis de la información tanto jurídica como espacial, con el fin de determinar a partir de método indirecto la ubicación, el área y los colindantes finales del predio objeto de estudio.', False),
    )

    _heading(doc, 'ANÁLISIS TÉCNICO', 2)

    _p_bold(doc,
        ('Se presenta en primera instancia el polígono que relaciona en la información geográfica de catastro; el cual cuenta con un área de: ', False),
        (_v(catastral, 'area_terreno_ha'), True),
        (' hectáreas, así mismo, se evidencia que su delimitación no es conforme con las coberturas presentes. Ahora, conociendo la ubicación del predio según la información de la malla catastral y de la plancha cartográfica XXXX a escala XXXX, se realiza un estudio multitemporal: ', False),
    )

    _tabla_dos_col(doc, [
        ('Banco de Imágenes de la ANT: ', ''),
        ('', ''),
        ('Imagen de SAS PLANET - Google Satélite', ''),
        ('', ''),
        ('Imagen Histórica Google: ', ''),
        ('', ''),
        ('Imagen Histórica Google: ', ''),
        ('', ''),
    ], h0='ANÁLISIS MULTITEMPORAL', h1='')
    _caption(doc, 'Tabla 4. Análisis multitemporal de coberturas')

    _p_bold(doc,
        ('El análisis multitemporal permite observar que para la ubicación relacionada a el predio objeto de estudio, las coberturas no han experimentado cambios significativos a lo largo del tiempo. ', False),
    )

    _p_bold(doc,
        ('En cuanto al relieve se determina a partir de un visor 3D que el área relacionada al predio objeto de estudio... ', False),
    )

    _caption(doc, '')
    _caption(doc, 'Ilustración 6. Visor 3D - Relieve')

    _caption(doc, '')
    _caption(doc, 'Ilustración 7. Polígono catastral y curvas de nivel')

    # ── 2.3 Análisis de colindantes — resumen ────────────
    _heading(doc, 'ANÁLISIS DE COLINDANTES', 2)

    _tabla_resumen_colindantes(doc, colindantes)
    _caption(doc, 'Tabla 5. Resumen de colindantes encontrados según insumos catastrales y registrales')

    # ── 2.4 Colindantes registrales ──────────────────────
    _p_bold(doc,
        ('La ', False),
        (_v(juridica, 'documento_origen_inmueble'), True),
        (', en la cual se adjudica el predio objeto de estudio señala la siguiente descripción de linderos:', False),
    )
    _caption(doc, '')
    _caption(doc, f"Ilustración 8. Linderos definidos en la {documento_origen_inmueble}")

    _p_bold(doc,
        ('La ', False),
        (_v(juridica, 'documento_ultima_anotacion'), True),
        (', corresponde a la última anotación registrada en la cadena de tradición del predio objeto de estudio que refleja al último propietario, allí se señala la siguiente descripción de linderos:', False),
    )
    _caption(doc, '')
    _caption(doc, f"Ilustración 9. Linderos definidos en la {documento_ultima_anotacion}")

    _p_bold(doc,
        ('Teniendo en cuenta lo anterior, junto con el análisis de los demás documentos asociados a los procesos traslaticios del predio, se evidencia que no hubo una actualización de los linderos del predio, manteniendo la descripción de la resolución de adjudicación. Igualmente se menciona que la descripción de la cabida y los linderos se hace como cuerpo cierto.', False),
    )
    _caption(doc, 'REEMPLAZAR')
    _caption(doc, 'Tabla 6. Resumen de colindantes mencionados en los insumos registrales')

    _p_bold(doc,
        ('En base a esta información, se realiza el análisis y comparativa entre los insumos catastrales y registrales, además de algunos predios de apoyo con el fin de asegurar la ubicación tanto de los colindantes como del predio objeto de estudio, esto permitió obtener: ', False),
    )

    _caption(doc, '')
    _caption(doc, 'Ilustración 10. Resumen de colindantes')

    _caption(doc, '')
    _caption(doc, 'Ilustración 11. Identificación de colindantes en malla catastral')

    _heading(doc, 'COLINDANTES SEGÚN INSUMOS REGISTRALES', 2)

    _tabla_colindantes(doc, col_registral,
        titulo_tabla='Tabla 7. Colindantes según insumos registrales')
    ilus_counter = _bloques_colindantes(doc, col_registral, ilustracion_inicio=12)

    # ── 2.5 Colindantes catastrales ──────────────────────
    _heading(doc, 'COLINDANTES SEGÚN INSUMOS CATASTRALES', 2)

    _p_bold(doc,
        ('De igual forma se toma la información de los insumos catastrales con el fin de identificar la posible existencia de colindantes que no hayan sido mencionados en la tradición jurídica del predio objeto de estudio:', False),
    )

    _tabla_colindantes(doc, col_catastral_lst,
        titulo_tabla='Tabla 8. Análisis de colindantes en malla catastral')
    ilus_counter = _bloques_colindantes(doc, col_catastral_lst, ilustracion_inicio=ilus_counter)


    _p_bold(doc,
        ('El anterior análisis de predios colindantes, junto con la exploración de predios de apoyo, nos permite realizar el proceso de reconstrucción de los documentos que dan origen a los colindantes y al predio objeto de estudio obteniendo:', False),
    )
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Polígonos de resoluciones de adjudicación de colindantes y predio objeto de estudio')
    _p_bold(doc,
        ('Al compararlo con la malla catastral vigente, se evidencia bastantes diferencias principalmente relacionadas a la forma de los predios y su ubicación. Por tal razón, para el proceso de validación de los colindantes se tuvo en cuenta las descripciones de cabida y linderos encontrados en las escrituras tanto del predio de estudio como los colindantes catastrales y registrales identificados.', False),
    )

    _p_bold(doc,
        ('Es importante señalar que el predio objeto de estudio ', False),
    )
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Comparativa polígonos de resoluciones de adjudicación y malla catastral vigente')
    _p_bold(doc,
        ('Entonces, aunque se tiene claro que se debe dar validez a las resoluciones más antiguas, también se debe tener en cuenta la Ortoimagen XXXXX y la cartografiá base disponible con la plancha catastral del IGAC XXXX del año XXXX, para validar linderos arcifinios, cambios de orientación y elementos materializados, pues, por ejemplo, en las resoluciones más antiguas, tomaban los vértices y trazaban líneas rectas, pero no tenían en cuenta estos cambios de dirección entre ellos. Los traslapes y vacíos entre resoluciones pueden deberse a inconsistencias en los levantamientos realizados por el INCORA, errores de calibración en los equipos utilizados, o a la temporalidad en la expedición de las resoluciones, pues entre algunas de ellas hay más de XX años de diferencia. ', False),
    )

    _p_bold(doc,
        ('Teniendo en cuenta lo anterior, y realizando la comparativa espacial junto con la revisión de linderos definidos en las escrituras más recientes, se ajustan los polígonos de los predios objeto de estudio y colindantes de la siguiente manera:', False),
    )
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Ajuste georreferenciación resoluciones de predio objeto de estudio y colindantes')

    _p_bold(doc,
        ('A partir del estudio de la información jurídica y catastral del predio objeto de estudio, junto con el Análisis Técnico y el Análisis de Colindantes, y realizado el ajuste del polígono del predio objeto de estudio y sus colindantes, se logra determinar un Área Final resultado del proceso de Gestión Catastral de ', False),
        (_v(predio, 'area_final_gc'), True),
        (', Hectáreas. En comparativa con las áreas obtenidas de otras fuentes, se observa:', False),
    )

    _p_bold(doc,
        ('', False),
    )
    _tabla_dos_col(doc, [
        ('Área VUR', _v(predio, 'area_vur', 'N/A') + ' HA'),
        ('Área Jurídica', _v(juridica, 'area_juridica_ha', 'N/A') + ' HA'),
        ('Área Polígono Malla Catastral', _v(catastral, 'area_terreno_ha', 'N/A') + ' HA'),
        ('Área R1/R2', _v(catastral, 'area_terreno_ha_r1r2', 'N/A') + ' HA'),
        ('ÁREA FINAL POLÍGONO GESTIÓN CATASTRAL', _v(predio, 'area_final_gc', 'N/A') + ' HA'),
    ], h0='CONCEPTO', h1='ÁREA')
    _caption(doc, 'Tabla 9. Comparativa de áreas')

    _p_bold(doc,
        ('Estas variaciones obedecen a ajustes derivados del proceso de georreferenciación, en el cual al momento de adaptar los planos de adjudicación a las coberturas a partir de los resultados obtenidos en el Análisis Técnico y Análisis de Colindantes, ciertos límites fueron modificados, dado que algunos colindantes también comparten tradición entre sí. En consecuencia, se reconoce como cabida real y técnica del predio objeto de estudio un área de ', False),
        (_v(predio, 'area_final_gc'), True),
        (' hectáreas.', False),
    )

    # ── 2.6 Colindantes finales (todos) ──────────────────
    _heading(doc, 'COLINDANTES FINALES LUEGO DE ANÁLISIS', 2)

    _p(doc, 'En conformidad con el ANÁLISIS DE COLINDANTES realizado, se especifica cuales colindantes mencionados de la resolución de adjudicación fueron verificados y cuales no, obteniendo como resultado:', False,)

    _tabla_colindantes_resumen_registrales(doc, col_registral,
        titulo_tabla='Tabla 10. Colindantes según insumos registrales')

    _p(doc, 'Finalmente se presenta el listado de colindantes finales, basado en los resultados del análisis realizado, permitiendo identificar cuáles colindantes requieren verificación en campo debido a la falta de información suficiente para validar su ubicación, así como cuáles fueron validados al coincidir con los colindantes jurídicos establecidos en la resolución de adjudicación. Para ello, se utilizaron como referencia los polígonos catastrales colindantes y la descripción de linderos contenida en dicha resolución.', False,)

    _tabla_colindantes_final(doc, list(colindantes),
                       titulo_tabla='Tabla 11. Colindantes finales')

    _caption(doc, '')
    _caption(doc, 'Ilustración X. Colindantes que requieren verificación en campo')

    _p(doc, 'Es relevante mencionar que los predios colindantes a los cuales no fue posible asociarle un FMI, se dejan como colindante final la persona en catastro, con el fin de que en campo se tenga una referencia y se pueda realizar la respectiva comprobación, ya que hasta este punto del análisis son considerados como presuntos baldíos luego de no poder comprobar la existencia de alguna tradición jurídica asociada a las personas allí mencionadas. ', False,)

    _heading(doc, 'OBSERVACIONES', 2)

    _p_bold(doc,
        ('XXXX', False),
    )

    _heading(doc, 'EXPEDIENTE ORFEO', 2)
    _p(doc, 'Se realiza la verificación de expedientes asociados al predio objeto de estudio, y se corrobora que actualmente este solo cuenta con el expediente XXXX, asociado a la Dirección de Acceso a Tierras, y el contenido corresponde al proceso de adquisición de predios', False,)
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Verificación de expediente ORFEO')

    _heading(doc, 'PROCESOS AGRARIOS', 2)
    _p(doc, 'Se realiza la verificación de la existencia de procesos agrarios asociados al predio objeto de estudio, y se corrobora que no cuenta con ninguno proceso en curso.', False,)
    _caption(doc, '')
    _caption(doc, 'Ilustración X. Revisión de Procesos Agrarios')

    # ── 2.7 Determinación del polígono final ─────────────
    _heading(doc, 'DETERMINACIÓN DEL POLÍGONO FINAL', 1)

    _heading(doc, 'UBICACIÓN', 2)

    _tabla_dos_col(doc, [
        ('Departamento (ubicación espacial)',      dpto),
        ('Municipio (ubicación espacial)',         mpio),
        ('Vereda / Corregimiento (espacial)',      vereda),
        ('¿Municipio diferente en VUR?',           _bool_s(localizacion.get('tiene_municipio_diferente_vur'))),
        ('Departamento (VUR)',                     _v(localizacion, 'dpto_vur', 'N/A')),
        ('Municipio (VUR)',                        _v(localizacion, 'municipio_vur', 'N/A')),
        ('Vereda / Corregimiento (VUR)',           _v(localizacion, 'vereda_vur', 'N/A')),
    ])
    _caption(doc, 'Tabla 12. Ubicación')

    _heading(doc, 'PRESENCIA DE CUERPOS DE AGUA', 2)

    _p(doc, 'Según la información cartográfica del IGAC (Plancha cartográfica XXXX), en el predio objeto de estudio se identifica ...', False,)

    _p(doc, 'Además, la verificación en los insumos espaciales como ortoimagen y cartografía base, evidenció la presencia de cuerpos de agua dentro del predio objeto de estudio, los cuales no se encontraban reflejados en la resolución de adjudicación. Se recomienda realizar una verificación en campo para determinar si estos cuerpos de agua son permanentes o temporales, y si cumplen con las características para ser considerados como humedales según la normatividad ambiental vigente.', False,)
    _caption(doc, 'Ilustración X. Cuerpos de agua')

    _heading(doc, 'FAJA PARALELA', 2)
    _p_bold(doc,
        ('De conformidad con la ', False),
        (_v(juridica, 'documento_origen_inmueble'), True),
        (', mediante la cual fue adjudicado el predio objeto de estudio, se debe tener presente que la adjudicación no comprende bienes de uso público, dentro de los cuales se incluyen las fuentes de agua y demás áreas legalmente protegidas.', False),
    )
    _p(doc, 'Adicionalmente, el artículo 677 del Código Civil Colombiano establece que los ríos y todas las aguas que corren por cauces naturales son bienes de la Unión de uso público en los respectivos territorios. En consecuencia, el área ocupada por dicho drenaje y las zonas asociadas a su protección deben excluirse del lindero jurídico del predio, por corresponder a bienes de uso público.', False),

    _p(doc, 'En atención a lo anterior, se realizó la delimitación de una faja paralela de treinta metros (30 m) sobre el cuerpo de agua identificado, obteniéndose un área de XX hectáreas, la cual corresponde al área objeto de exclusión por su condición de protección hídrica y uso público.', False),

    _p(doc, 'No se identifican que los cuerpos de agua presentes en el predio objeto de estudio tengan un ancho mayor a 30 metros, por lo cual, y según lo definido en el Decreto 2245 de 2017, Decreto 2811 de 1974 y Decreto 1210 de 2020, no aplica la Faja Paralela.', False),

    _caption(doc, '')
    _caption(doc, 'Ilustración X. Faja paralela')

    _tabla_dos_col(doc, [
        ('Área Faja Paralela', 'XX HA + XXXX M2'),
        ('Área Final Polígono Gestión Catastral', _v(predio, 'area_final_gc', 'N/A') + ' HA + XXXX M2'),
        ('Área Sin Restricción Ambiental', 'XX HA + XXXX M2'),
    ], h0='CONCEPTO', h1='ÁREA (HA))')
    _caption(doc, 'Tabla 13. Área de faja paralela, área sin restricción ambiental y área a adjudicar')


    # ═══════════════════════════════════════════════════════
    # 3. EVALUACIÓN FINAL
    # ═══════════════════════════════════════════════════════
    _heading(doc, 'EVALUACIÓN FINAL', 1)


    _p_bold(doc,
        ('El análisis técnico y jurídico del predio denominado ', False),
        (nombre_actual, True),
        (', identificado con número predial ', False),
        (num_predial, True),
        (' y FMI ', False),
        (fmi, True),
        (' registra un área jurídica de ', False),
        (_v(juridica, 'area_juridica_ha', 'N/A'), True),
        (' ha de acuerdo con la ', False),
        (_v(juridica, 'documento_origen_inmueble'), True),
        ('.', False),
        ('. El análisis jurídico, catastral y técnico permitió establecer los linderos, áreas y demás características del bien inmueble objeto de estudio.', False),
    )

    _p(doc,
        'Una vez realizado el análisis de colindantes, la georreferenciación y materialización de los planos, descripciones y/o carteras obtenidas, a partir de método XXXX, se pudo determinar un área final del predio de XX HA + XXXX M2, la cual difiere del área jurídica en un X.XX%, con una diferencia de XX HA + XXXX M2.')

    _p(doc,
        'La obtención de esta área a partir de método XXXX implico un análisis de colindantes que consistió en la revisión y georreferenciación de los planos asociados al predio objeto de estudio así como de predios colindantes o cercanos al mismo, se identifican las características más relevantes con el fin de encontrar coincidencias puntuales que permitan validar la delimitación de cada plano o cada polígono recreado, de igual forma una revisión de la tradición jurídica de cada predio con el fin de validar las colindancias en los procesos traslaticios que ha tenido el área de estudio. Para este proceso también se tuvo en cuenta un análisis técnico con base en la cartografía base, así como un estudio multitemporal de las coberturas existentes en el área de estudio, lo cual permitió validar la ubicación de los colindantes y del predio objeto de estudio, así como también validar la existencia de linderos arcifinios y/o artificiales que permitan establecer límites más precisos. Es importante mencionar que el método XXXX es un proceso de validación indirecta, por lo cual se recomienda realizar una verificación en campo para validar los resultados obtenidos a partir de este método, así como para validar la presencia de cuerpos de agua y la delimitación de la faja paralela, con el fin de asegurar que el área determinada como resultado del proceso de Gestión Catastral corresponda con las condiciones reales del predio objeto de estudio.', False,)

    _p(doc,
        'Asimismo, se determinó un área objeto de exclusión por condición de protección hídrica y uso público de XX HA + XXXX M2, por lo cual se obtuvo un área final para compra de XX HA + XXXX M2.')

    _tabla_evaluacion_final_resumen(doc)
    _caption(doc, 'Tabla 14. Evaluación Final')

    _p(doc,
        'CONCEPTO DE ESCRITURACIÓN')

    _p(doc,
        'Las diferencias encontradas obedecen a la ausencia de planos protocolizados y a la descripción vaga de cabida que permitan verificar con mayor precisión la delimitación del predio. Por lo tanto, los resultados aquí consignados tienen carácter aproximado y no pueden asumirse como determinación técnica definitiva.')

    _p(doc, 'Se recomienda realizar un levantamiento predial planimétrico como condición indispensable para validar el área y los linderos, y para sustentar cualquier proceso de actualización catastral, registral o de gestión administrativa sobre el inmueble. Este concepto se emite como dictamen técnico de carácter informativo y no constituye acto administrativo ni declaración de validez sobre la titularidad de derechos reales.')

    # ═══════════════════════════════════════════════════════
    # 4. CRUCE CON DETERMINANTES
    # ═══════════════════════════════════════════════════════
    _heading(doc, 'CRUCE CON DETERMINANTES', 1)
    _caption(doc, 'CRUCE TOOLBOX')
    _caption(doc, 'CONDICIONANTE 1: X 0.0 HA 0%')
    _caption(doc, 'CONDICIONANTE 2: X 0.0 HA 0%')
    _caption(doc, 'CONDICIONANTE 3: X 0.0 HA 0%')
    _caption(doc, 'RESTRICCIÓN 1: X 0.0 HA 0%')
    _caption(doc, 'RESTRICCIÓN 2: X 0.0 HA 0%')
    _caption(doc, 'RESTRICCIÓN 3: X 0.0 HA 0%')
    _caption(doc, 'CRUCE EXPERIENCE')
    _caption(doc, 'PROCEDE 1: X 0.0 HA 0%')
    _caption(doc, 'PROCEDE 2: X 0.0 HA 0%')
    _caption(doc, 'PROCEDE 3: X 0.0 HA 0%')

    _heading(doc, 'DE ACUERDO CON LA REVISIÓN TECNICA Y CATASTRAL, SE CONSIDERA PERTINENTE CONTINUAR CON LOS TRAMITES DE COMPRA DEL PREDIO.', 1)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer